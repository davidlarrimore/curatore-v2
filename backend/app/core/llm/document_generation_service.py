# backend/app/services/document_generation_service.py
"""
Document Generation Service - Creates PDF, DOCX, and CSV files.

This is a compartmentalized service designed to be replaceable with Carbone
or another document generation solution in the future.

Current Implementation:
    - PDF: WeasyPrint (markdown -> HTML -> PDF)
    - DOCX: python-docx (markdown -> DOCX with basic formatting)
    - CSV: Python csv module (dict list -> CSV bytes)

Future:
    Replace this entire service with Carbone client for template-based
    document generation with better formatting control.

Usage:
    from app.core.llm.document_generation_service import document_generation_service

    # Generate PDF from markdown
    pdf_bytes = await document_generation_service.generate_pdf(
        content="# Hello World\n\nThis is a test.",
        title="My Document"
    )

    # Generate DOCX from markdown
    docx_bytes = await document_generation_service.generate_docx(
        content="# Report\n\n- Item 1\n- Item 2",
        title="Report"
    )

    # Generate CSV from data
    csv_bytes = await document_generation_service.generate_csv(
        data=[{"name": "Alice", "age": 30}, {"name": "Bob", "age": 25}]
    )
"""

from __future__ import annotations

import csv
import io
import logging
import re
from typing import Any, Dict, List, Optional

import markdown

logger = logging.getLogger("curatore.document_generation")


# Default CSS for PDF generation
DEFAULT_PDF_CSS = """
@page {
    size: letter;
    margin: 1in;
}

body {
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto,
                 "Helvetica Neue", Arial, sans-serif;
    font-size: 11pt;
    line-height: 1.6;
    color: #333;
}

h1 {
    font-size: 24pt;
    font-weight: 600;
    color: #1a1a1a;
    margin-top: 0;
    margin-bottom: 16pt;
    border-bottom: 2px solid #4f46e5;
    padding-bottom: 8pt;
}

h2 {
    font-size: 18pt;
    font-weight: 600;
    color: #1a1a1a;
    margin-top: 24pt;
    margin-bottom: 12pt;
}

h3 {
    font-size: 14pt;
    font-weight: 600;
    color: #333;
    margin-top: 18pt;
    margin-bottom: 8pt;
}

p {
    margin-bottom: 12pt;
}

ul, ol {
    margin-bottom: 12pt;
    padding-left: 24pt;
}

li {
    margin-bottom: 4pt;
}

code {
    font-family: "SF Mono", Monaco, "Courier New", monospace;
    font-size: 10pt;
    background-color: #f5f5f5;
    padding: 2pt 4pt;
    border-radius: 3pt;
}

pre {
    font-family: "SF Mono", Monaco, "Courier New", monospace;
    font-size: 10pt;
    background-color: #f5f5f5;
    padding: 12pt;
    border-radius: 4pt;
    overflow-x: auto;
    margin-bottom: 12pt;
}

blockquote {
    border-left: 4px solid #4f46e5;
    padding-left: 16pt;
    margin-left: 0;
    color: #555;
    font-style: italic;
}

table {
    width: 100%;
    border-collapse: collapse;
    margin-bottom: 12pt;
}

th, td {
    border: 1px solid #ddd;
    padding: 8pt;
    text-align: left;
}

th {
    background-color: #f5f5f5;
    font-weight: 600;
}

tr:nth-child(even) {
    background-color: #fafafa;
}

a {
    color: #4f46e5;
    text-decoration: none;
}

hr {
    border: none;
    border-top: 1px solid #ddd;
    margin: 24pt 0;
}

.title-page {
    text-align: center;
    padding-top: 200pt;
}

.title-page h1 {
    border-bottom: none;
    font-size: 32pt;
}

.metadata {
    color: #666;
    font-size: 10pt;
    margin-bottom: 24pt;
}
"""


class DocumentGenerationService:
    """
    Service for generating documents (PDF, DOCX, CSV) from content.

    This service is designed to be compartmentalized and replaceable.
    The interface is intentionally simple to allow easy swapping with
    Carbone or another document generation backend in the future.
    """

    def __init__(self):
        """Initialize the document generation service."""
        self._markdown_converter = markdown.Markdown(
            extensions=[
                "tables",
                "fenced_code",
                "toc",
                "nl2br",
            ]
        )

    # =========================================================================
    # PDF GENERATION
    # =========================================================================

    async def generate_pdf(
        self,
        content: str,
        title: Optional[str] = None,
        css: Optional[str] = None,
        include_title_page: bool = False,
    ) -> bytes:
        """
        Generate a PDF from markdown content.

        Args:
            content: Markdown content to convert
            title: Optional document title (shown in title page if enabled)
            css: Optional custom CSS (uses default if not provided)
            include_title_page: If True, adds a title page at the beginning

        Returns:
            PDF file as bytes

        Raises:
            RuntimeError: If PDF generation fails
        """
        try:
            # Import WeasyPrint lazily to avoid import errors if not installed
            from weasyprint import CSS, HTML

            # Convert markdown to HTML
            self._markdown_converter.reset()
            html_content = self._markdown_converter.convert(content)

            # Build full HTML document
            css_content = css or DEFAULT_PDF_CSS

            title_page_html = ""
            if include_title_page and title:
                title_page_html = f"""
                <div class="title-page">
                    <h1>{self._escape_html(title)}</h1>
                </div>
                <div style="page-break-after: always;"></div>
                """

            full_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="utf-8">
                <title>{self._escape_html(title or 'Document')}</title>
            </head>
            <body>
                {title_page_html}
                {html_content}
            </body>
            </html>
            """

            # Generate PDF
            html_doc = HTML(string=full_html)
            css_doc = CSS(string=css_content)
            pdf_bytes = html_doc.write_pdf(stylesheets=[css_doc])

            logger.debug(f"Generated PDF: {len(pdf_bytes)} bytes")
            return pdf_bytes

        except ImportError as e:
            logger.error(f"WeasyPrint not installed: {e}")
            raise RuntimeError(
                "PDF generation requires WeasyPrint. "
                "Install with: pip install weasyprint"
            ) from e
        except Exception as e:
            logger.error(f"PDF generation failed: {e}")
            raise RuntimeError(f"PDF generation failed: {e}") from e

    # =========================================================================
    # DOCX GENERATION
    # =========================================================================

    async def generate_docx(
        self,
        content: str,
        title: Optional[str] = None,
    ) -> bytes:
        """
        Generate a DOCX file from markdown content.

        Supports basic markdown formatting:
        - Headers (h1-h3)
        - Bold and italic text
        - Unordered and ordered lists
        - Paragraphs

        Args:
            content: Markdown content to convert
            title: Optional document title (added as document title property)

        Returns:
            DOCX file as bytes

        Raises:
            RuntimeError: If DOCX generation fails
        """
        try:
            # Import python-docx lazily
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Inches, Pt

            doc = Document()

            # Set document title property
            if title:
                doc.core_properties.title = title

            # Parse and add content
            lines = content.split("\n")
            current_list_type = None  # 'ul' or 'ol'
            list_counter = 0

            for line in lines:
                stripped = line.strip()

                if not stripped:
                    # Empty line - end any current list
                    current_list_type = None
                    list_counter = 0
                    continue

                # Headers
                if stripped.startswith("### "):
                    p = doc.add_heading(stripped[4:], level=3)
                    current_list_type = None
                elif stripped.startswith("## "):
                    p = doc.add_heading(stripped[3:], level=2)
                    current_list_type = None
                elif stripped.startswith("# "):
                    p = doc.add_heading(stripped[2:], level=1)
                    current_list_type = None
                # Unordered list
                elif stripped.startswith("- ") or stripped.startswith("* "):
                    text = stripped[2:]
                    p = doc.add_paragraph(style="List Bullet")
                    self._add_formatted_text(p, text)
                    current_list_type = "ul"
                # Ordered list
                elif re.match(r"^\d+\.\s", stripped):
                    text = re.sub(r"^\d+\.\s", "", stripped)
                    p = doc.add_paragraph(style="List Number")
                    self._add_formatted_text(p, text)
                    current_list_type = "ol"
                    list_counter += 1
                # Horizontal rule
                elif stripped in ("---", "***", "___"):
                    # Add a paragraph with a bottom border to simulate HR
                    p = doc.add_paragraph()
                    current_list_type = None
                # Regular paragraph
                else:
                    p = doc.add_paragraph()
                    self._add_formatted_text(p, stripped)
                    current_list_type = None

            # Save to bytes
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            docx_bytes = buffer.read()

            logger.debug(f"Generated DOCX: {len(docx_bytes)} bytes")
            return docx_bytes

        except ImportError as e:
            logger.error(f"python-docx not installed: {e}")
            raise RuntimeError(
                "DOCX generation requires python-docx. "
                "Install with: pip install python-docx"
            ) from e
        except Exception as e:
            logger.error(f"DOCX generation failed: {e}")
            raise RuntimeError(f"DOCX generation failed: {e}") from e

    def _add_formatted_text(self, paragraph, text: str):
        """
        Add text to a paragraph with basic markdown formatting (bold, italic).

        Args:
            paragraph: python-docx Paragraph object
            text: Text with potential markdown formatting
        """
        # Pattern to match **bold**, *italic*, and ***bold italic***
        pattern = r"(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))"

        for match in re.finditer(pattern, text):
            if match.group(2):  # ***bold italic***
                run = paragraph.add_run(match.group(2))
                run.bold = True
                run.italic = True
            elif match.group(3):  # **bold**
                run = paragraph.add_run(match.group(3))
                run.bold = True
            elif match.group(4):  # *italic*
                run = paragraph.add_run(match.group(4))
                run.italic = True
            elif match.group(5):  # plain text
                paragraph.add_run(match.group(5))

    # =========================================================================
    # CSV GENERATION
    # =========================================================================

    async def generate_csv(
        self,
        data: List[Dict[str, Any]],
        columns: Optional[List[str]] = None,
        include_bom: bool = True,
    ) -> bytes:
        """
        Generate a CSV file from a list of dictionaries.

        Args:
            data: List of dictionaries (each dict is a row)
            columns: Optional list of column names (auto-detected if not provided)
            include_bom: If True, includes UTF-8 BOM for Excel compatibility

        Returns:
            CSV file as bytes (UTF-8 encoded)

        Raises:
            RuntimeError: If CSV generation fails
            ValueError: If data is empty and no columns provided
        """
        try:
            if not data and not columns:
                raise ValueError("Cannot generate CSV: no data or columns provided")

            # Auto-detect columns from data if not provided
            if columns is None:
                if data:
                    # Collect all unique keys from all rows
                    all_keys = set()
                    for row in data:
                        all_keys.update(row.keys())
                    columns = sorted(all_keys)
                else:
                    columns = []

            # Build CSV
            buffer = io.StringIO()
            writer = csv.DictWriter(
                buffer,
                fieldnames=columns,
                extrasaction="ignore",  # Ignore extra keys in rows
            )

            writer.writeheader()
            for row in data:
                writer.writerow(row)

            csv_content = buffer.getvalue()

            # Encode to bytes with optional BOM
            if include_bom:
                csv_bytes = b"\xef\xbb\xbf" + csv_content.encode("utf-8")
            else:
                csv_bytes = csv_content.encode("utf-8")

            logger.debug(f"Generated CSV: {len(csv_bytes)} bytes, {len(data)} rows")
            return csv_bytes

        except Exception as e:
            logger.error(f"CSV generation failed: {e}")
            raise RuntimeError(f"CSV generation failed: {e}") from e

    # =========================================================================
    # UTILITY METHODS
    # =========================================================================

    def _escape_html(self, text: str) -> str:
        """Escape HTML special characters."""
        return (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
            .replace("'", "&#39;")
        )

    def get_supported_formats(self) -> List[str]:
        """Get list of supported output formats."""
        return ["pdf", "docx", "csv"]

    async def generate(
        self,
        content: str,
        format: str,
        title: Optional[str] = None,
        data: Optional[List[Dict[str, Any]]] = None,
        **kwargs,
    ) -> bytes:
        """
        Universal generation method - routes to the appropriate generator.

        Args:
            content: Markdown content (for pdf/docx)
            format: Output format (pdf, docx, csv)
            title: Optional document title
            data: Data for CSV generation (overrides content)
            **kwargs: Additional format-specific options

        Returns:
            Generated document as bytes

        Raises:
            ValueError: If format is not supported
            RuntimeError: If generation fails
        """
        format_lower = format.lower().strip()

        if format_lower == "pdf":
            return await self.generate_pdf(
                content=content,
                title=title,
                css=kwargs.get("css"),
                include_title_page=kwargs.get("include_title_page", False),
            )
        elif format_lower == "docx":
            return await self.generate_docx(
                content=content,
                title=title,
            )
        elif format_lower == "csv":
            if data is None:
                raise ValueError("CSV generation requires 'data' parameter")
            return await self.generate_csv(
                data=data,
                columns=kwargs.get("columns"),
                include_bom=kwargs.get("include_bom", True),
            )
        else:
            supported = ", ".join(self.get_supported_formats())
            raise ValueError(
                f"Unsupported format '{format}'. Supported formats: {supported}"
            )


# Global service instance
document_generation_service = DocumentGenerationService()
